# Python
from django.shortcuts import render, redirect
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.urls import reverse_lazy
from django.views.generic import ListView, CreateView, UpdateView, DeleteView, DetailView, TemplateView
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.shortcuts import get_object_or_404
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os
from resume_builder.models import (
    WorkExperience, Education, Project, Certification, 
    Award, Language, TechnicalSkill, PersonalInformation, Resume
)
from resume_builder.forms import (
    WorkExperienceForm, EducationForm, ProjectForm, CertificationForm,
    AwardForm, LanguageForm, TechnicalSkillForm, PersonalInformationForm, ResumeForm
)

# Personal Information CRUD Views
class PersonalInformationListView(LoginRequiredMixin, ListView):
    model = PersonalInformation
    template_name = 'resume_builder/personal_information/personal_information_list.html'
    context_object_name = 'personal_information_list'

    def get_queryset(self):
        return PersonalInformation.objects.filter(resume__user=self.request.user)

class PersonalInformationCreateView(LoginRequiredMixin, CreateView):
    model = PersonalInformation
    form_class = PersonalInformationForm
    template_name = 'resume_builder/personal_information/personal_information_form.html'
    success_url = reverse_lazy('resume_builder_web:personal_information_list')

    def form_valid(self, form):
        # Assign the first resume belonging to the user
        resume = Resume.objects.filter(user=self.request.user).first()
        if not resume:
            form.add_error(None, 'You must create a resume before adding personal information.')
            return self.form_invalid(form)
        form.instance.resume = resume
        response = super().form_valid(form)
        from django.contrib import messages
        messages.success(self.request, 'Personal information saved successfully!')
        return response

class PersonalInformationUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = PersonalInformation
    form_class = PersonalInformationForm
    template_name = 'resume_builder/personal_information/personal_information_form.html'
    success_url = reverse_lazy('resume_builder_web:personal_information_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class PersonalInformationDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = PersonalInformation
    template_name = 'resume_builder/personal_information/personal_information_confirm_delete.html'
    success_url = reverse_lazy('resume_builder_web:personal_information_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class PersonalInformationDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = PersonalInformation
    template_name = 'resume_builder/personal_information/personal_information_detail.html'
    context_object_name = 'personal_information'

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class WorkExperienceListView(LoginRequiredMixin, ListView):
    model = WorkExperience
    template_name = 'resume_builder/work_experience/work_experience_list.html'
    context_object_name = 'experiences'

    def get_queryset(self):
        return WorkExperience.objects.filter(resume__user=self.request.user)

class WorkExperienceCreateView(LoginRequiredMixin, CreateView):
    model = WorkExperience
    form_class = WorkExperienceForm
    template_name = 'resume_builder/work_experience/work_experience_form.html'
    success_url = reverse_lazy('resume_builder_web:work_experience_list')

    def form_valid(self, form):
        # Ensure the resume belongs to the user
        resume = form.cleaned_data['resume']
        if resume.user != self.request.user:
            form.add_error('resume', 'You do not own this resume.')
            return self.form_invalid(form)
        return super().form_valid(form)

class WorkExperienceUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = WorkExperience
    form_class = WorkExperienceForm
    template_name = 'resume_builder/work_experience/work_experience_form.html'
    success_url = reverse_lazy('resume_builder_web:work_experience_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class WorkExperienceDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = WorkExperience
    template_name = 'resume_builder/work_experience/work_experience_confirm_delete.html'
    success_url = reverse_lazy('resume_builder_web:work_experience_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class WorkExperienceDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = WorkExperience
    template_name = 'resume_builder/work_experience/work_experience_detail.html'
    context_object_name = 'experience'

    def test_func(self):
        return self.get_object().resume.user == self.request.user


# Education CRUD Views
class EducationListView(LoginRequiredMixin, ListView):
    model = Education
    template_name = 'resume_builder/education/education_list.html'
    context_object_name = 'educations'

    def get_queryset(self):
        return Education.objects.filter(resume__user=self.request.user)

class EducationCreateView(LoginRequiredMixin, CreateView):
    model = Education
    form_class = EducationForm
    template_name = 'resume_builder/education/education_form.html'
    success_url = reverse_lazy('resume_builder_web:education_list')

    def form_valid(self, form):
        resume = form.cleaned_data['resume']
        if resume.user != self.request.user:
            form.add_error('resume', 'You do not own this resume.')
            return self.form_invalid(form)
        return super().form_valid(form)

class EducationUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = Education
    form_class = EducationForm
    template_name = 'resume_builder/education/education_form.html'
    success_url = reverse_lazy('resume_builder_web:education_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class EducationDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Education
    template_name = 'resume_builder/education/education_confirm_delete.html'
    success_url = reverse_lazy('resume_builder_web:education_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class EducationDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = Education
    template_name = 'resume_builder/education/education_detail.html'
    context_object_name = 'education'

    def test_func(self):
        return self.get_object().resume.user == self.request.user


# Project CRUD Views
class ProjectListView(LoginRequiredMixin, ListView):
    model = Project
    template_name = 'resume_builder/project/project_list.html'
    context_object_name = 'projects'

    def get_queryset(self):
        return Project.objects.filter(resume__user=self.request.user)

class ProjectCreateView(LoginRequiredMixin, CreateView):
    model = Project
    form_class = ProjectForm
    template_name = 'resume_builder/project/project_form.html'
    success_url = reverse_lazy('resume_builder_web:project_list')

    def form_valid(self, form):
        resume = form.cleaned_data['resume']
        if resume.user != self.request.user:
            form.add_error('resume', 'You do not own this resume.')
            return self.form_invalid(form)
        return super().form_valid(form)

class ProjectUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = Project
    form_class = ProjectForm
    template_name = 'resume_builder/project/project_form.html'
    success_url = reverse_lazy('resume_builder_web:project_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class ProjectDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Project
    template_name = 'resume_builder/project/project_confirm_delete.html'
    success_url = reverse_lazy('resume_builder_web:project_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class ProjectDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = Project
    template_name = 'resume_builder/project/project_detail.html'
    context_object_name = 'project'

    def test_func(self):
        return self.get_object().resume.user == self.request.user


# Certification CRUD Views
class CertificationListView(LoginRequiredMixin, ListView):
    model = Certification
    template_name = 'resume_builder/certification/certification_list.html'
    context_object_name = 'certifications'

    def get_queryset(self):
        return Certification.objects.filter(resume__user=self.request.user)

class CertificationCreateView(LoginRequiredMixin, CreateView):
    model = Certification
    form_class = CertificationForm
    template_name = 'resume_builder/certification/certification_form.html'
    success_url = reverse_lazy('resume_builder_web:certification_list')

    def form_valid(self, form):
        resume = form.cleaned_data['resume']
        if resume.user != self.request.user:
            form.add_error('resume', 'You do not own this resume.')
            return self.form_invalid(form)
        return super().form_valid(form)

class CertificationUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = Certification
    form_class = CertificationForm
    template_name = 'resume_builder/certification/certification_form.html'
    success_url = reverse_lazy('resume_builder_web:certification_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class CertificationDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Certification
    template_name = 'resume_builder/certification/certification_confirm_delete.html'
    success_url = reverse_lazy('resume_builder_web:certification_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class CertificationDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = Certification
    template_name = 'resume_builder/certification/certification_detail.html'
    context_object_name = 'certification'

    def test_func(self):
        return self.get_object().resume.user == self.request.user


# Award CRUD Views
class AwardListView(LoginRequiredMixin, ListView):
    model = Award
    template_name = 'resume_builder/award/award_list.html'
    context_object_name = 'awards'

    def get_queryset(self):
        return Award.objects.filter(resume__user=self.request.user)

class AwardCreateView(LoginRequiredMixin, CreateView):
    model = Award
    form_class = AwardForm
    template_name = 'resume_builder/award/award_form.html'
    success_url = reverse_lazy('resume_builder_web:award_list')

    def form_valid(self, form):
        resume = form.cleaned_data['resume']
        if resume.user != self.request.user:
            form.add_error('resume', 'You do not own this resume.')
            return self.form_invalid(form)
        return super().form_valid(form)

class AwardUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = Award
    form_class = AwardForm
    template_name = 'resume_builder/award/award_form.html'
    success_url = reverse_lazy('resume_builder_web:award_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class AwardDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Award
    template_name = 'resume_builder/award/award_confirm_delete.html'
    success_url = reverse_lazy('resume_builder_web:award_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class AwardDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = Award
    template_name = 'resume_builder/award/award_detail.html'
    context_object_name = 'award'

    def test_func(self):
        return self.get_object().resume.user == self.request.user


# Language CRUD Views
class LanguageListView(LoginRequiredMixin, ListView):
    model = Language
    template_name = 'resume_builder/language/language_list.html'
    context_object_name = 'languages'

    def get_queryset(self):
        return Language.objects.filter(resume__user=self.request.user)

class LanguageCreateView(LoginRequiredMixin, CreateView):
    model = Language
    form_class = LanguageForm
    template_name = 'resume_builder/language/language_form.html'
    success_url = reverse_lazy('resume_builder_web:language_list')

    def form_valid(self, form):
        resume = form.cleaned_data['resume']
        if resume.user != self.request.user:
            form.add_error('resume', 'You do not own this resume.')
            return self.form_invalid(form)
        return super().form_valid(form)

class LanguageUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = Language
    form_class = LanguageForm
    template_name = 'resume_builder/language/language_form.html'
    success_url = reverse_lazy('resume_builder_web:language_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class LanguageDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Language
    template_name = 'resume_builder/language/language_confirm_delete.html'
    success_url = reverse_lazy('resume_builder_web:language_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class LanguageDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = Language
    template_name = 'resume_builder/language/language_detail.html'
    context_object_name = 'language'

    def test_func(self):
        return self.get_object().resume.user == self.request.user


# TechnicalSkill CRUD Views
class TechnicalSkillListView(LoginRequiredMixin, ListView):
    model = TechnicalSkill
    template_name = 'resume_builder/technical_skill/technical_skill_list.html'
    context_object_name = 'technical_skills'

    def get_queryset(self):
        return TechnicalSkill.objects.filter(resume__user=self.request.user)

class TechnicalSkillCreateView(LoginRequiredMixin, CreateView):
    model = TechnicalSkill
    form_class = TechnicalSkillForm
    template_name = 'resume_builder/technical_skill/technical_skill_form.html'
    success_url = reverse_lazy('resume_builder_web:technical_skill_list')

    def form_valid(self, form):
        resume = form.cleaned_data['resume']
        if resume.user != self.request.user:
            form.add_error('resume', 'You do not own this resume.')
            return self.form_invalid(form)
        return super().form_valid(form)

class TechnicalSkillUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = TechnicalSkill
    form_class = TechnicalSkillForm
    template_name = 'resume_builder/technical_skill/technical_skill_form.html'
    success_url = reverse_lazy('resume_builder_web:technical_skill_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class TechnicalSkillDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = TechnicalSkill
    template_name = 'resume_builder/technical_skill/technical_skill_confirm_delete.html'
    success_url = reverse_lazy('resume_builder_web:technical_skill_list')

    def test_func(self):
        return self.get_object().resume.user == self.request.user

class TechnicalSkillDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = TechnicalSkill
    template_name = 'resume_builder/technical_skill/technical_skill_detail.html'
    context_object_name = 'technical_skill'

    def test_func(self):
        return self.get_object().resume.user == self.request.user


# Resume Preview Views
class ResumePreviewView(LoginRequiredMixin, TemplateView):
    """Base class for resume preview views"""
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        resume_id = self.kwargs.get('resume_id')
        resume = get_object_or_404(Resume, id=resume_id, user=self.request.user)
        
        # Get all resume data
        context['resume'] = resume
        context['personal_info'] = resume.personal_information.first()  # Get personal information
        context['work_experiences'] = resume.work_experiences.all()
        context['educations'] = resume.educations.all()
        context['projects'] = resume.projects.all()
        context['technical_skills'] = resume.technical_skills.all()
        context['certifications'] = resume.certifications.all()
        context['awards'] = resume.awards.all()
        context['languages'] = resume.languages.all()
        return context


class ClassicResumePreviewView(ResumePreviewView):
    template_name = 'resume_builder/preview/template_classic.html'


class ModernResumePreviewView(ResumePreviewView):
    template_name = 'resume_builder/preview/template_modern.html'


class CreativeResumePreviewView(ResumePreviewView):
    template_name = 'resume_builder/preview/template_creative.html'


class TechnicalResumePreviewView(ResumePreviewView):
    template_name = 'resume_builder/preview/template_technical.html'


class MinimalResumePreviewView(ResumePreviewView):
    template_name = 'resume_builder/preview/template_minimal.html'


class ResumeListView(LoginRequiredMixin, ListView):
    model = Resume
    template_name = 'resume_builder/resume_list.html'
    context_object_name = 'resumes'

    def get_queryset(self):
        return Resume.objects.filter(user=self.request.user)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['user_resume'] = Resume.objects.filter(user=self.request.user).first()
        return context


class ResumeCreateView(LoginRequiredMixin, CreateView):
    model = Resume
    form_class = ResumeForm
    template_name = 'resume_builder/resume_form.html'
    success_url = reverse_lazy('resume_builder_web:resume_list')

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)


class ResumeDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = Resume
    template_name = 'resume_builder/resume_detail.html'
    context_object_name = 'resume'

    def test_func(self):
        return self.get_object().user == self.request.user

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        resume = self.get_object()
        
        # Get all resume data
        context.update({
            'personal_info': PersonalInformation.objects.filter(resume=resume).first(),
            'work_experiences': WorkExperience.objects.filter(resume=resume).order_by('-end_date'),
            'educations': Education.objects.filter(resume=resume).order_by('-end_date'),
            'projects': Project.objects.filter(resume=resume).order_by('-start_date'),
            'certifications': Certification.objects.filter(resume=resume).order_by('-issue_date'),
            'awards': Award.objects.filter(resume=resume).order_by('-issue_date'),
            'languages': Language.objects.filter(resume=resume),
            'technical_skills': TechnicalSkill.objects.filter(resume=resume),
        })
        
        return context


class ResumeDownloadPDFView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = Resume
    template_name = 'resume_builder/resume_pdf_template.html'
    context_object_name = 'resume'

    def test_func(self):
        return self.get_object().user == self.request.user

    def get(self, request, *args, **kwargs):
        resume = self.get_object()
        
        # Get all resume data
        context = {
            'resume': resume,
            'personal_info': PersonalInformation.objects.filter(resume=resume).first(),
            'work_experiences': WorkExperience.objects.filter(resume=resume).order_by('-end_date'),
            'educations': Education.objects.filter(resume=resume).order_by('-end_date'),
            'projects': Project.objects.filter(resume=resume).order_by('-start_date'),
            'certifications': Certification.objects.filter(resume=resume).order_by('-issue_date'),
            'awards': Award.objects.filter(resume=resume).order_by('-issue_date'),
            'languages': Language.objects.filter(resume=resume),
            'technical_skills': TechnicalSkill.objects.filter(resume=resume),
        }
        
        # Render the template
        html_string = render_to_string(self.template_name, context)
        
        # Create PDF
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{resume.title}_resume.pdf"'
        
        # Generate PDF
        pisa_status = pisa.CreatePDF(html_string, dest=response)
        
        if pisa_status.err:
            return HttpResponse('PDF generation failed', status=500)
        
        return response


class ResumeDownloadDOCXView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    model = Resume
    context_object_name = 'resume'

    def test_func(self):
        return self.get_object().user == self.request.user

    def get(self, request, *args, **kwargs):
        resume = self.get_object()
        
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading(resume.title, 0)
        title.alignment = 1  # Center alignment
        
        # Get all resume data
        personal_info = PersonalInformation.objects.filter(resume=resume).first()
        work_experiences = WorkExperience.objects.filter(resume=resume).order_by('-end_date')
        educations = Education.objects.filter(resume=resume).order_by('-end_date')
        projects = Project.objects.filter(resume=resume).order_by('-start_date')
        certifications = Certification.objects.filter(resume=resume).order_by('-issue_date')
        awards = Award.objects.filter(resume=resume).order_by('-issue_date')
        languages = Language.objects.filter(resume=resume)
        technical_skills = TechnicalSkill.objects.filter(resume=resume)
        
        # Personal Information
        if personal_info:
            doc.add_heading('Personal Information', level=1)
            p = doc.add_paragraph()
            p.add_run(f'Name: {personal_info.full_name}\n')
            p.add_run(f'Email: {personal_info.email}\n')
            if personal_info.phone:
                p.add_run(f'Phone: {personal_info.phone}\n')
            if personal_info.address:
                p.add_run(f'Address: {personal_info.address}\n')
            if personal_info.linkedin_url:
                p.add_run(f'LinkedIn: {personal_info.linkedin_url}\n')
            if personal_info.github_url:
                p.add_run(f'GitHub: {personal_info.github_url}\n')
        
        # Summary
        if resume.summary:
            doc.add_heading('Summary', level=1)
            doc.add_paragraph(resume.summary)
        
        # Work Experience
        if work_experiences:
            doc.add_heading('Work Experience', level=1)
            for exp in work_experiences:
                p = doc.add_paragraph()
                p.add_run(f'{exp.job_title} at {exp.company}').bold = True
                p.add_run(f' ({exp.start_date} - {exp.end_date or "Present"})')
                doc.add_paragraph(exp.description)
        
        # Education
        if educations:
            doc.add_heading('Education', level=1)
            for edu in educations:
                p = doc.add_paragraph()
                p.add_run(f'{edu.degree} from {edu.institution}').bold = True
                p.add_run(f' ({edu.start_date} - {edu.end_date})')
                if edu.gpa:
                    p.add_run(f' GPA: {edu.gpa}')
        
        # Technical Skills
        if technical_skills:
            doc.add_heading('Technical Skills', level=1)
            for skill in technical_skills:
                p = doc.add_paragraph()
                p.add_run(f'{skill.technology.name}').bold = True
                p.add_run(f' - {skill.get_proficiency_display()}')
        
        # Projects
        if projects:
            doc.add_heading('Projects', level=1)
            for project in projects:
                p = doc.add_paragraph()
                p.add_run(f'{project.title}').bold = True
                p.add_run(f' ({project.start_date} - {project.end_date or "Present"})')
                doc.add_paragraph(project.description)
                if project.url:
                    doc.add_paragraph(f'URL: {project.url}')
        
        # Certifications
        if certifications:
            doc.add_heading('Certifications', level=1)
            for cert in certifications:
                p = doc.add_paragraph()
                p.add_run(f'{cert.name}').bold = True
                p.add_run(f' - {cert.issuer} ({cert.issue_date})')
                if cert.expiration_date:
                    p.add_run(f' - Expires: {cert.expiration_date}')
        
        # Awards
        if awards:
            doc.add_heading('Awards', level=1)
            for award in awards:
                p = doc.add_paragraph()
                p.add_run(f'{award.title}').bold = True
                p.add_run(f' - {award.issuer} ({award.issue_date})')
                doc.add_paragraph(award.description)
        
        # Languages
        if languages:
            doc.add_heading('Languages', level=1)
            for lang in languages:
                p = doc.add_paragraph()
                p.add_run(f'{lang.name}').bold = True
                p.add_run(f' - {lang.get_proficiency_display()}')
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create response
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{resume.title}_resume.docx"'
        
        return response
